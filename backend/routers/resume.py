from fastapi import APIRouter, HTTPException, UploadFile, File, Depends
from services.credits import charge_action
from fastapi.responses import StreamingResponse, Response as FileResponse
from pydantic import BaseModel
from agents.resume_builder_agent import build_from_qa, build_from_linkedin, enhance_bullet, stream_build, edit_resume_with_instruction
from agents.resume_adaptor_agent import adapt_resume, parse_job_description, generate_cover_letter, score_resume_vs_jd
from agents.resume_scorer_agent import score_resume
from services.linkedin_scraper import enrich_linkedin_input
import json
import io
import httpx
from bs4 import BeautifulSoup

router = APIRouter()


class ScoreRequest(BaseModel):
    resume: dict
    target_role: str = ""


@router.post("/score")
async def score_resume_endpoint(req: ScoreRequest):
    """Comprehensive resume score — 7 dimensions, 100 points. Free for all users."""
    result = await score_resume(req.resume, req.target_role)
    return result

# System prompt for structured extraction
SYSTEM_EXTRACT = """You are a resume data extraction specialist. Your ONLY job is to read an existing resume document and extract the data into structured JSON.

CRITICAL RULES:
1. EXTRACT ONLY — never invent, infer, or add information not explicitly written in the document
2. Preserve the EXACT job titles, company names, and dates as written
3. Preserve the EXACT industry and domain — never change it
4. Copy bullet points verbatim or very close to verbatim
5. If a field is not present, leave it as empty string or empty array
6. The "title" field = their most recent or primary job title exactly as written
7. For LinkedIn PDF exports: the format has sections like "Experience", "Education", "Skills" — extract all of them

Output ONLY valid JSON:
{
  "personal": {"name":"","email":"","phone":"","location":"","linkedin":"","github":"","website":"","title":""},
  "summary": "",
  "experience": [{"company":"","role":"","start":"","end":"","location":"","current":false,"bullets":[]}],
  "education": [{"institution":"","degree":"","field":"","start":"","end":"","gpa":""}],
  "skills": {"technical":[],"soft":[],"languages":[],"certifications":[]},
  "projects": [{"name":"","description":"","tech":[],"link":"","bullets":[]}],
  "achievements": [],
  "volunteer": []
}"""


class LinkedInRequest(BaseModel):
    linkedin_text: str


class QABuildRequest(BaseModel):
    conversation: list[dict]


class BulletRequest(BaseModel):
    bullet: str
    role: str
    company: str


class AdaptRequest(BaseModel):
    resume: dict
    jd_text: str
    company_name: str = ""
    role_name: str = ""


class CoverLetterRequest(BaseModel):
    resume: dict
    jd_text: str
    tone: str = "professional"


class AdaptScoreRequest(BaseModel):
    resume: dict
    jd_text: str


class FetchJDRequest(BaseModel):
    url: str

class ParseFileRequest(BaseModel):
    file_text: str
    file_name: str = ""


@router.post("/build/linkedin", dependencies=[Depends(charge_action("resume_build"))])
async def build_from_linkedin_route(req: LinkedInRequest):
    try:
        # Enrich URL with scraped public data + name extraction before sending to Claude
        enriched = await enrich_linkedin_input(req.linkedin_text)
        resume = await build_from_linkedin(enriched)
        return {"resume": resume}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/build/qa", dependencies=[Depends(charge_action("resume_build"))])
async def build_from_qa_route(req: QABuildRequest):
    try:
        resume = await build_from_qa(req.conversation)
        return {"resume": resume}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/enhance-bullet", dependencies=[Depends(charge_action("chat_message"))])
async def enhance_bullet_route(req: BulletRequest):
    try:
        enhanced = await enhance_bullet(req.bullet, req.role, req.company)
        return {"enhanced": enhanced}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/adapt", dependencies=[Depends(charge_action("resume_adapt"))])
async def adapt_resume_route(req: AdaptRequest):
    try:
        jd_parsed = await parse_job_description(req.jd_text)
        result = await adapt_resume(req.resume, req.jd_text, jd_parsed, company_name=req.company_name, role_name=req.role_name)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/cover-letter", dependencies=[Depends(charge_action("cover_letter"))])
async def cover_letter_route(req: CoverLetterRequest):
    try:
        letter = await generate_cover_letter(req.resume, req.jd_text, req.tone)
        return {"cover_letter": letter}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/score-vs-jd")
async def score_route(req: AdaptScoreRequest):
    try:
        score = await score_resume_vs_jd(req.resume, req.jd_text)
        return score
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _repair_truncated_json(text: str) -> str:
    """Close any open strings/arrays/objects left by a truncated Claude response."""
    stack = []
    in_string = False
    escape_next = False
    for ch in text:
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch in ('{', '['):
            stack.append(ch)
        elif ch == '}' and stack and stack[-1] == '{':
            stack.pop()
        elif ch == ']' and stack and stack[-1] == '[':
            stack.pop()
    closer = ('"' if in_string else '')
    for ch in reversed(stack):
        closer += '}' if ch == '{' else ']'
    return text + closer


def _parse_json_resilient(raw: str, label: str = "") -> dict:
    """Parse JSON; if truncated, attempt bracket repair before failing."""
    import json as _json
    from loguru import logger
    try:
        return _json.loads(raw)
    except _json.JSONDecodeError:
        try:
            repaired = _repair_truncated_json(raw)
            result = _json.loads(repaired)
            if label:
                logger.warning(f"Repaired truncated JSON for {label}")
            return result
        except Exception:
            raise


@router.post("/parse-file", dependencies=[Depends(charge_action("resume_build"))])
async def parse_file_route(req: ParseFileRequest):
    """Parse resume text — used when text is pre-extracted client-side."""
    from services.claude_service import complete_claude_json

    prompt = f"""File: {req.file_name}

Document content:
{req.file_text[:12000]}

Extract every field exactly as written. Do not change job titles, industries, or add information."""

    try:
        raw = await complete_claude_json(SYSTEM_EXTRACT, [{"role": "user", "content": prompt}], max_tokens=8192)
        resume = _parse_json_resilient(raw, req.file_name)
        return {"resume": resume}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload", dependencies=[Depends(charge_action("resume_build"))])
async def upload_resume_file(file: UploadFile = File(...)):
    """
    Upload a PDF, DOCX, or TXT file as binary — properly extracts text using
    PyMuPDF (for PDF) or python-docx (for DOCX) before passing to Claude.
    This is the CORRECT way to handle PDF uploads, not readAsText().
    """
    from services.claude_service import complete_claude_json as ccj

    filename = file.filename or "resume"
    ext = filename.lower().rsplit(".", 1)[-1]
    content = await file.read()

    text = ""

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text("text") + "\n\n"
            doc.close()
            text = text.strip()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not read PDF: {str(e)}. Make sure the PDF is not a scanned image.")

    elif ext == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            seen = set()
            paragraphs = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t and t not in seen:
                    seen.add(t)
                    paragraphs.append(t)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in seen:
                            seen.add(t)
                            paragraphs.append(t)
            text = "\n".join(paragraphs)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Could not read DOCX: {str(e)}")

    else:
        # Plain text
        text = content.decode("utf-8", errors="ignore")

    # Strip null bytes and other control chars that break Claude's JSON output
    import re as _re
    text = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', ' ', text).strip()
    # Collapse runs of whitespace/blank lines
    text = _re.sub(r'\n{3,}', '\n\n', text)

    if not text or len(text) < 30:
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from this file. For scanned PDFs, please copy-paste the text instead."
        )

    prompt = f"""File name: {filename}
File type: {ext.upper()}

Document content:
{text[:18000]}

Extract every field exactly as it appears. Do not change job titles, industries, or add information not in the document."""

    try:
        raw = await ccj(SYSTEM_EXTRACT, [{"role": "user", "content": prompt}], max_tokens=8192)
        resume = _parse_json_resilient(raw, filename)
        return {"resume": resume, "chars_extracted": len(text), "pages": text.count("\n\n")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI parsing failed: {str(e)}")


class ResumeEditRequest(BaseModel):
    instruction: str
    current_resume: dict


@router.post("/edit", dependencies=[Depends(charge_action("chat_message"))])
async def edit_resume_route(req: ResumeEditRequest):
    try:
        updated = await edit_resume_with_instruction(req.instruction, req.current_resume)
        return {"resume": updated}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _extract_jd_text_from_html(html: str) -> tuple[str, str]:
    """
    Extract job description text from HTML.
    Returns (text, title).
    Looks for JSON-LD structured data, meta tags, and article/main/section content.
    """
    import re as _re

    soup = BeautifulSoup(html, "lxml")

    # Extract title before removing anything
    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    # Prefer og:title for cleaner job title
    og_title = soup.find("meta", property="og:title")
    if og_title and og_title.get("content"):
        title = og_title["content"].strip()

    parts: list[str] = []

    # ── 1. JSON-LD structured data (JobPosting schema) ───────────────────────
    for script in soup.find_all("script", type="application/ld+json"):
        try:
            data = json.loads(script.string or "")
            if isinstance(data, list):
                data = data[0] if data else {}
            jtype = data.get("@type", "")
            if "job" in str(jtype).lower() or "Job" in str(jtype):
                if data.get("title"):
                    parts.append(f"Job Title: {data['title']}")
                if data.get("hiringOrganization"):
                    org = data["hiringOrganization"]
                    if isinstance(org, dict):
                        parts.append(f"Company: {org.get('name', '')}")
                if data.get("jobLocation"):
                    loc = data["jobLocation"]
                    if isinstance(loc, dict):
                        addr = loc.get("address", {})
                        if isinstance(addr, dict):
                            parts.append(f"Location: {addr.get('addressLocality', '')} {addr.get('addressRegion', '')} {addr.get('addressCountry', '')}".strip())
                if data.get("description"):
                    parts.append(f"Description:\n{data['description']}")
                if data.get("qualifications"):
                    parts.append(f"Qualifications:\n{data['qualifications']}")
                if data.get("responsibilities"):
                    parts.append(f"Responsibilities:\n{data['responsibilities']}")
                if data.get("skills"):
                    parts.append(f"Skills: {data['skills']}")
                if data.get("baseSalary"):
                    parts.append(f"Salary: {json.dumps(data['baseSalary'])}")
        except Exception:
            pass

    # ── 2. Meta description ───────────────────────────────────────────────────
    og_desc = soup.find("meta", property="og:description")
    if og_desc and og_desc.get("content"):
        parts.append(f"Summary: {og_desc['content']}")
    meta_desc = soup.find("meta", attrs={"name": "description"})
    if meta_desc and meta_desc.get("content"):
        parts.append(f"Meta: {meta_desc['content']}")

    # ── 3. Visible content from job-specific selectors ───────────────────────
    # Remove boilerplate before extracting text
    for tag in soup(["script", "style", "nav", "footer", "header", "noscript", "aside", "meta", "link"]):
        tag.decompose()

    main_content = None
    for selector in [
        "[class*='job-description']",
        "[class*='jobDescription']",
        "[class*='job_description']",
        "[id*='job-description']",
        "[id*='jobDescription']",
        "[class*='description__text']",
        "[class*='show-more-less-html']",
        "[class*='jobDetail']",
        "[class*='job-detail']",
        "[class*='posting']",
        "article",
        "main",
        "[role='main']",
        "section",
    ]:
        found = soup.select_one(selector)
        if found:
            txt = found.get_text(separator="\n", strip=True)
            if len(txt) > 200:
                main_content = found
                break

    if main_content:
        parts.append(main_content.get_text(separator="\n", strip=True))
    elif not parts:
        body = soup.find("body")
        if body:
            parts.append(body.get_text(separator="\n", strip=True))

    raw_text = "\n\n".join(parts)

    # Clean up excessive blank lines
    lines = [line.strip() for line in raw_text.splitlines()]
    cleaned_lines: list[str] = []
    prev_blank = False
    for line in lines:
        if line:
            cleaned_lines.append(line)
            prev_blank = False
        elif not prev_blank:
            cleaned_lines.append("")
            prev_blank = True

    text = "\n".join(cleaned_lines).strip()
    if len(text) > 8000:
        text = text[:8000] + "\n\n[Content truncated...]"

    return text, title


async def _fetch_jd_with_playwright(url: str) -> tuple[str, str]:
    """Render JS-heavy pages (Deloitte, Naukri, LinkedIn jobs) with Playwright."""
    try:
        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=True,
                args=[
                    "--disable-blink-features=AutomationControlled",
                    "--no-sandbox",
                    "--disable-setuid-sandbox",
                    "--disable-dev-shm-usage",
                    "--disable-gpu",
                    "--no-first-run",
                ],
            )
            context = await browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
                viewport={"width": 1440, "height": 900},
                locale="en-US",
                extra_http_headers={
                    "Accept-Language": "en-US,en;q=0.9",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                },
            )
            await context.add_init_script(
                "Object.defineProperty(navigator, 'webdriver', { get: () => undefined });"
            )
            page = await context.new_page()
            try:
                await page.goto(url, wait_until="domcontentloaded", timeout=25000)
                await page.wait_for_timeout(3000)
                html = await page.content()
            except Exception:
                html = await page.content()
            await browser.close()
            return _extract_jd_text_from_html(html)
    except ImportError:
        return "", ""
    except Exception as e:
        from loguru import logger
        logger.warning(f"Playwright JD fetch failed: {e}")
        return "", ""


@router.post("/fetch-jd")
async def fetch_jd_route(req: FetchJDRequest):
    """Fetch a job description or LinkedIn profile from a URL.
    Strategy:
    1. Try httpx first (fast for simple pages)
    2. If content <500 meaningful chars, fall back to Playwright headless browser
    """
    url = req.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="Invalid URL — must start with http:// or https://")

    # Route LinkedIn profile URLs through the dedicated LinkedIn scraper
    if "linkedin.com/in/" in url:
        from services.linkedin_scraper import fetch_profile_text
        text = await fetch_profile_text(url)
        if text and len(text) > 50:
            return {"text": text, "title": "LinkedIn Profile", "source": "linkedin_scraper"}
        raise HTTPException(status_code=422, detail="Could not extract profile data from LinkedIn. Please paste your profile content manually.")

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
    }

    text = ""
    title = ""
    used_playwright = False

    # ── Strategy 1: httpx (fast) ─────────────────────────────────────────────
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=15) as client:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            text, title = _extract_jd_text_from_html(response.text)
    except httpx.TimeoutException:
        pass  # Will try Playwright below
    except httpx.HTTPStatusError as e:
        if e.response.status_code in (403, 429, 503):
            pass  # Likely bot-blocked — try Playwright
        else:
            raise HTTPException(status_code=502, detail=f"Remote server returned {e.response.status_code}.")
    except Exception:
        pass  # Try Playwright

    # ── Strategy 2: Playwright fallback for JS-rendered pages ────────────────
    # Deloitte, Naukri, LinkedIn jobs pages require JS rendering
    needs_playwright = (
        not text
        or len(text.replace("\n", "").replace(" ", "")) < 500
    )
    js_heavy_domains = ["deloitte.com", "naukri.com", "linkedin.com/jobs", "workday.com", "greenhouse.io", "lever.co"]
    if needs_playwright or any(d in url for d in js_heavy_domains):
        from loguru import logger
        logger.info(f"Falling back to Playwright for {url} (httpx text length: {len(text)})")
        pw_text, pw_title = await _fetch_jd_with_playwright(url)
        if pw_text and len(pw_text.replace("\n", "").replace(" ", "")) > len(text.replace("\n", "").replace(" ", "")):
            text = pw_text
            title = pw_title or title
            used_playwright = True

    if not text:
        raise HTTPException(status_code=422, detail="Could not extract job description from this URL. The page may require login or is heavily JS-rendered. Please paste the job description text manually.")

    return {"text": text, "title": title, "used_playwright": used_playwright}


class ExportPDFRequest(BaseModel):
    html: str
    name: str = "resume"


@router.post("/export-pdf", dependencies=[Depends(charge_action("pdf_download"))])
async def export_pdf_endpoint(req: ExportPDFRequest):
    """Render resume HTML to PDF using Playwright headless Chromium."""
    try:
        from playwright.async_api import async_playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=True,
                args=[
                    "--no-sandbox",
                    "--disable-setuid-sandbox",
                    "--disable-dev-shm-usage",
                    "--disable-accelerated-2d-canvas",
                    "--disable-gpu",
                ],
            )
            page = await browser.new_page()
            await page.set_content(req.html, wait_until="networkidle")
            await page.wait_for_timeout(800)
            pdf_bytes = await page.pdf(
                format="A4",
                print_background=True,
                margin={"top": "0mm", "right": "0mm", "bottom": "0mm", "left": "0mm"},
            )
            await browser.close()

        safe_name = "".join(c for c in req.name if c.isalnum() or c in " _-").strip() or "resume"
        return FileResponse(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}_resume.pdf"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")
