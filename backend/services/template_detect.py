"""
Detect which Mithra template best matches an uploaded resume's visual design.

An arbitrary PDF's exact layout can't be reproduced, but its *design language* can
be read from the file: typeface family, accent colours, column structure and
background blocks. We map those signals to the closest built-in template so an
adapted resume comes back looking like the one the user uploaded.

Zero API cost — pure local analysis with PyMuPDF / python-docx.
"""
from collections import Counter
from typing import Optional

TEMPLATES = ("modern", "classic", "minimal", "bold", "tech")

_SERIF_HINTS = ("times", "georgia", "garamond", "serif", "cambria", "book", "minion", "baskerville", "palatino")
_MONO_HINTS = ("courier", "mono", "consolas", "menlo")


def _rgb(color_int: int) -> tuple[int, int, int]:
    return ((color_int >> 16) & 255, (color_int >> 8) & 255, color_int & 255)


def _is_greyish(r: int, g: int, b: int) -> bool:
    return max(abs(r - g), abs(g - b), abs(r - b)) < 28


def detect_template_from_pdf(content: bytes) -> dict:
    """Returns {template, confidence, signals} — never raises."""
    signals: dict = {}
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        page = doc[0]
        pw, ph = page.rect.width, page.rect.height

        fonts: Counter = Counter()
        accent_colors: Counter = Counter()
        x_starts: list[float] = []

        data = page.get_text("dict")
        for block in data.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    txt = (span.get("text") or "").strip()
                    if not txt:
                        continue
                    fonts[(span.get("font") or "").lower()] += len(txt)
                    r, g, b = _rgb(span.get("color", 0))
                    # A saturated, non-grey colour = an accent
                    if not _is_greyish(r, g, b):
                        accent_colors[(r, g, b)] += len(txt)
                    x_starts.append(span.get("bbox", [0])[0])

        # ── Typeface family ──────────────────────────────────────────────
        serif_wt = sum(w for f, w in fonts.items() if any(h in f for h in _SERIF_HINTS))
        mono_wt = sum(w for f, w in fonts.items() if any(h in f for h in _MONO_HINTS))
        total_wt = max(sum(fonts.values()), 1)
        signals["serif_ratio"] = round(serif_wt / total_wt, 2)
        signals["mono_ratio"] = round(mono_wt / total_wt, 2)

        # ── Accent colour presence ───────────────────────────────────────
        accent_wt = sum(accent_colors.values())
        signals["accent_ratio"] = round(accent_wt / total_wt, 2)
        top_accent = accent_colors.most_common(1)[0][0] if accent_colors else None
        signals["accent_rgb"] = top_accent

        # ── Two-column / sidebar detection ───────────────────────────────
        # A sidebar shows up as a dense cluster of text starting in the left third
        # AND another cluster past the middle.
        left = sum(1 for x in x_starts if x < pw * 0.33)
        right = sum(1 for x in x_starts if x > pw * 0.42)
        two_col = bool(left > 12 and right > 12 and (left / max(len(x_starts), 1)) > 0.22)
        signals["two_column"] = two_col

        # ── Large filled background block (dark sidebar / header band) ───
        dark_block = False
        try:
            for d in page.get_drawings():
                fill = d.get("fill")
                if not fill:
                    continue
                r, g, b = [int(c * 255) for c in fill[:3]]
                rect = d.get("rect")
                if rect is None:
                    continue
                area = (rect.x1 - rect.x0) * (rect.y1 - rect.y0)
                if area > (pw * ph) * 0.10:  # covers >10% of the page
                    lum = 0.299 * r + 0.587 * g + 0.114 * b
                    if lum < 130:
                        dark_block = True
                    if not _is_greyish(r, g, b):
                        accent_colors[(r, g, b)] += 400  # weight coloured bands heavily
        except Exception:
            pass
        signals["dark_block"] = dark_block
        doc.close()

        template = _decide(signals)
        return {"template": template, "confidence": _confidence(signals, template), "signals": signals}
    except Exception as e:
        return {"template": "modern", "confidence": 0.0, "signals": {"error": str(e)[:120]}}


def detect_template_from_docx(content: bytes) -> dict:
    """DOCX gives us fonts and colours; columns are rarer, so signals are simpler."""
    signals: dict = {}
    try:
        import io
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        fonts: Counter = Counter()
        accents = 0
        total = 0
        for para in doc.paragraphs:
            for run in para.runs:
                t = (run.text or "").strip()
                if not t:
                    continue
                total += len(t)
                fname = (run.font.name or "").lower()
                if fname:
                    fonts[fname] += len(t)
                col = run.font.color
                try:
                    if col and col.rgb is not None:
                        r, g, b = col.rgb[0], col.rgb[1], col.rgb[2]
                        if not _is_greyish(r, g, b):
                            accents += len(t)
                except Exception:
                    pass
        total = max(total, 1)
        serif_wt = sum(w for f, w in fonts.items() if any(h in f for h in _SERIF_HINTS))
        mono_wt = sum(w for f, w in fonts.items() if any(h in f for h in _MONO_HINTS))
        signals = {
            "serif_ratio": round(serif_wt / total, 2),
            "mono_ratio": round(mono_wt / total, 2),
            "accent_ratio": round(accents / total, 2),
            "two_column": len(doc.tables) > 0,   # 2-col DOCX resumes are usually tables
            "dark_block": False,
            "accent_rgb": None,
        }
        template = _decide(signals)
        return {"template": template, "confidence": _confidence(signals, template), "signals": signals}
    except Exception as e:
        return {"template": "modern", "confidence": 0.0, "signals": {"error": str(e)[:120]}}


def _decide(s: dict) -> str:
    """Map visual signals → closest Mithra template."""
    serif = s.get("serif_ratio", 0)
    mono = s.get("mono_ratio", 0)
    accent = s.get("accent_ratio", 0)
    two_col = s.get("two_column", False)
    dark = s.get("dark_block", False)
    rgb = s.get("accent_rgb")

    # Monospace or a dark panel → the terminal-styled "tech" template
    if mono > 0.15 or dark:
        return "tech"
    # Serif body copy → the traditional "classic" template
    if serif > 0.35:
        return "classic"
    # Two-column with colour → "tech" (has a sidebar); without colour → modern
    if two_col and accent > 0.03:
        return "tech"
    # A strong warm/red accent → "bold"
    if rgb and accent > 0.05:
        r, g, b = rgb
        if r > 150 and r > g + 40 and r > b + 40:
            return "bold"
    # Colour present but restrained → "modern"; almost no colour → "minimal"
    if accent >= 0.02:
        return "modern"
    return "minimal"


def _confidence(s: dict, template: str) -> float:
    if s.get("error"):
        return 0.0
    if template == "tech" and (s.get("mono_ratio", 0) > 0.15 or s.get("dark_block")):
        return 0.85
    if template == "classic" and s.get("serif_ratio", 0) > 0.5:
        return 0.85
    if template == "minimal" and s.get("accent_ratio", 1) < 0.01:
        return 0.75
    return 0.6


def detect_template(content: bytes, ext: str) -> dict:
    ext = (ext or "").lower()
    if ext == "pdf":
        return detect_template_from_pdf(content)
    if ext == "docx":
        return detect_template_from_docx(content)
    return {"template": "modern", "confidence": 0.0, "signals": {"note": "plain text — no design to detect"}}
